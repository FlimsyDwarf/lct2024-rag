from functools import reduce

from docx import Document
from docx.shared import Pt, Inches

from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser


def make_single_query(main_topic, block_name, query=None, keyword="блок"):
    """
    Создает полный единичный запрос поиска для модели

    Args:
        main_topic: главная тема отчета
        block_name: название одного блока
        query: запрос
        keyword: ключевое слово разбиения на блоки

    Returns: полный запрос поиска для модели

    """

    block_name = block_name[len(keyword):].strip()
    if query is not None:
        return f"{main_topic}: {block_name}: {query}"
    return f"{main_topic}: {block_name}"


def get_info_from_template(template_file, keyword) -> tuple[str, str, dict[str, list[str]]]:
    """
    Обрабатывает шаблон пользователя

    Args:
        template_file: путь к шаблону
        keyword: ключевое слово, по которому просиходит разбиение на блоки

    Returns: heading: str - заголовок шаблона
             main_topic: str - главная тема отчета
             blocks: dict[str, list[str]] - разбиение файла на блоки в виде словаря:
             описание блока(первая его строка) -> все его подпункты одной строкой

    """

    template = Document(template_file)

    # Сохраняем заголовок шаблона - такой же заголовок будет у отчета
    heading = template.paragraphs[0].text.strip()

    # Сохраняем главную тему отчета(например, "Компани 'Северсталь'") для использования в запросах
    main_topic = template.paragraphs[1].text.strip()

    # Собираем блоки из подзапросов по порядку. Один вопрос - один подпункт
    blocks = dict()

    block = []
    block_name = ""
    sub_query = ""
    for paragraph in template.paragraphs[2:]:

        # Разделяем шаблон на блоки по ключевому слову
        if keyword in paragraph.text.lower().strip():
            if sub_query:
                block.append(sub_query.strip())
            sub_query = ""

            if block:
                blocks[block_name] = block
            block_name = paragraph.text.strip()
            block = []

        elif block_name:
            # Внутри блока берем запросы по пунктам/нумерации, созданной пользователем: один пункт/номер - один запрос
            if paragraph._element.xpath('./w:pPr/w:numPr'):  # Проверка, что параграф является нумерацией или пунктом
                if sub_query:
                    block.append(sub_query.strip())
                sub_query = ""
            sub_query += paragraph.text + '\n'

    block.append(sub_query.strip())
    blocks[block_name] = block

    for name, vals in blocks.items():
        print(name)
        print('\t', end='')
        print(*vals, sep='\n\t')

    return heading, main_topic, blocks


def create_final_report(output_file, heading, main_topic, reports: dict[str, dict[str, dict[str, str]]], keyword):
    """
    Создает конечный файл, в котором будет сохранен результат работы программы

    Args:
        output_file: путь к конечному файлу, в котором будет сохранен результат
        heading: заголовок шаблона
        main_topic: главная тема отчета
        reports: словарь сгенерированных ответов от модели на каждый подпункт каждого блока
        keyword: ключевое слово разбиения на блоки

    """

    final_report = Document()

    style = final_report.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Добавляем заголовок отчета
    heading = final_report.add_paragraph(heading)
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].bold = True
    heading.alignment = 1

    # Добавляем основную тему отчета
    heading = final_report.add_paragraph(main_topic)
    heading.runs[0].font.size = Pt(12)
    heading.alignment = 1

    for block_name, block_reports in reports.items():

        # Добавляем название блока в отчет
        paragraph = final_report.add_paragraph()
        runner = paragraph.add_run(block_name)
        runner.bold = True

        for ind, (sub_query, sub_report) in enumerate(block_reports.items(), start=1):
            # Добавляем название пользовательского запроса в отчет с 1 отступом
            paragraph = final_report.add_paragraph()
            runner = paragraph.add_run(f"\t{ind}) {sub_query}")
            runner.bold = True

            # Суммаризуем ответы и ссылки модели
            report_info, report_sources = __summarise_model_output(
                main_topic, block_name, sub_report, sub_query, keyword)

            # Добавляем ответ на ind-ый запрос в блоке в отчет с 1 отступом и ещё чуть-чуть(1 отступ = 0.5 inch)
            paragraph = final_report.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.7)
            paragraph.add_run(f"{report_info}\n")

            # Добавляем источники на ind-ый запрос в блоке в отчет с 1 отступом и ещё чуть-чуть(1 отступ = 0.5 inch)
            runner = paragraph.add_run(report_sources)
            runner.font.size = Pt(10)
            runner.italic = True

    final_report.save(output_file)


def __summarise_model_output(main_topic, block_name, sub_report, sub_query, keyword):
    """
        Суммаризует ответы модели на один пользовательский запрос
    """

    sub_reports_info = ""
    sub_reports_sources = ""

    # Собираем ответы в один большой и все ссылки в одну строку
    for extra_query, extra_report in sub_report.items():
        try:
            sub_report_info, sub_report_sources = extra_report.rsplit('\n\n', 1)
            sub_reports_info += sub_report_info + '\n'
            sub_reports_sources += sub_report_sources + '\n'
        except ValueError:
            print(f"Информация не найдена: {sub_query} \n\n{extra_query} \n{extra_report} \n\n")

    if sub_reports_info:
        # Суммаризуем гптхой все ответы в нужной теме по нужному запросу пользователя
        report_info = summarise_info().invoke(
            {"sub_reports": sub_reports_info,
             "query": sub_query,
             "topic": make_single_query(main_topic, block_name, keyword=keyword)})[
            "report"].strip()
        if sub_reports_sources:
            # Скалдываем гптхой все источники в одну строку с разделением через пробел
            # И приводим в надлежащий читаемый вид эти источники
            report_sources = (("Источники:\n" +
                               reduce(lambda x, y: x + '- ' + y + '\n', [''] + list(
                                   summarise_sources().invoke({"sub_sources": sub_reports_sources})[
                                       "sources"].split())))
                              .strip())
        else:
            report_sources = "Источники отсутствуют"
    else:
        report_info = "Информация не найдена"
        report_sources = "Источники отсутствуют"

    return report_info, report_sources


def summarise_info():
    """
    Возвращает модель с нужным промптом для суммаризации ответов на один пользовательский вопрос
    """

    # LLM
    llm = ChatOpenAI(
        model="gpt-4o",
        temperature=0,
        max_tokens=None,
        timeout=None,
        max_retries=2,
    )

    # Prompt
    divide_prompt = PromptTemplate(
        template="""Вы профессионально отвечаете на заданный запрос в заданной теме, используя только данный вам текст.
        Вам на вход подаётся текст, тема и запрос.
        Посмотрите на исходный текст и тему, составьте точный, однозначный, 
        не содержащий противоречий и повторений информации,
        логичный ответ на заданный запрос, учитывая заданную тему, используя только и только данный вам текст.
        Напишите только ответ на поставленный запрос.\n
        Ответ дай в виде JSON в одну строку с одним ключом 'report' и ничего больше. Без символов переноса строк. \n
        Вот текст: {sub_reports}.\n 
        Вот тема: {topic}.\n 
        Вот запрос: {query}.\n 
        Твой JSON:\n """,
        input_variables=["sub_reports", "query", "topic"],
    )

    question_divide = divide_prompt | llm | JsonOutputParser()
    return question_divide


def summarise_sources():
    """
    Возвращает модель с нужным промптом для поиска всех ссылок в тексте
    """

    # LLM
    llm = ChatOpenAI(
        model="gpt-4o",
        temperature=0,
        max_tokens=None,
        timeout=None,
        max_retries=2,
    )

    # Prompt
    divide_prompt = PromptTemplate(
        template="""Вы профессионал в поиске ссылок на интернет-ресурсы любого формата в заданном тексте.
        Вам на вход подаётся текст с большим количеством ссылок.
        Посмотрите на исходный текст, найдите каждую ссылку.
        Запишите все найденные ссылки, которые есть в исходном тексте, без изменений через пробел.
        Идентичные ссылки напишите один раз.\n
        Ответ дай в виде JSON в одну строку с одним ключом 'sources' и значением обязательно строкового типа и ничего больше.
        Без символов переноса строк. \n
        Вот текст: {sub_sources}.\n 
        Твой JSON:\n """,
        input_variables=["sub_sources"],
    )

    question_divide = divide_prompt | llm | JsonOutputParser()
    return question_divide
